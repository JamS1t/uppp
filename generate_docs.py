"""
generate_docs.py
Generates uppp_documentation.docx — ITE 206L Final Project Documentation
Run: python generate_docs.py
Output: uppp_documentation.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

FONT_NAME = "Book Antiqua"
FONT_SIZE = Pt(12)
CODE_FONT = "Courier New"
CODE_SIZE = Pt(10)

doc = Document()

# ── Page setup: Letter 8.5" x 11" ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.5)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)


# ── Helper functions ────────────────────────────────────────────────────────

def set_run_font(run, bold=False, italic=False, size=FONT_SIZE, font=FONT_NAME):
    run.font.name = font
    run.font.size = size
    run.bold      = bold
    run.italic    = italic
    # Force font for East Asian / complex scripts too
    r = run._r
    rPr = r.get_or_add_rPr()
    for tag in (qn("w:rFonts"),):
        el = rPr.find(tag)
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:ascii"), font)
        el.set(qn("w:hAnsi"), font)


def set_para_spacing(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_before=Pt(0), space_after=Pt(6)):
    fmt = para.paragraph_format
    fmt.alignment          = align
    fmt.line_spacing_rule  = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before       = space_before
    fmt.space_after        = space_after


def add_para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
             italic=False, space_after=Pt(6)):
    p = doc.add_paragraph()
    set_para_spacing(p, align=align, space_after=space_after)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return p


def add_chapter_title(text):
    """Bold, centered chapter heading."""
    p = doc.add_paragraph()
    set_para_spacing(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=Pt(12), space_after=Pt(12))
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_section_heading(text):
    """Bold, left-aligned sub-heading."""
    p = doc.add_paragraph()
    set_para_spacing(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=Pt(10), space_after=Pt(4))
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_sub_heading(text):
    """Bold, italic, left-aligned third-level heading."""
    p = doc.add_paragraph()
    set_para_spacing(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=Pt(8), space_after=Pt(4))
    run = p.add_run(text)
    set_run_font(run, bold=True, italic=True)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_before=Pt(0), space_after=Pt(3))
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_code_block(lines):
    """Add a shaded code block."""
    for line in lines:
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.alignment         = WD_ALIGN_PARAGRAPH.LEFT
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.space_before      = Pt(0)
        fmt.space_after       = Pt(0)
        fmt.left_indent       = Inches(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
        # Light grey shading
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:highlight")
        shd.set(qn("w:val"), "none")
        pPr = p._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"), "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd2)
    # small gap after block
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(4)
    gap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def page_break():
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════

for _ in range(6):
    add_para("", align=WD_ALIGN_PARAGRAPH.CENTER)

add_para("UPPP: A Community-Based Link Sharing Platform",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=Pt(24))

add_para("A Final Project in partial fulfillment of the requirements for the course",
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=Pt(4))
add_para("ITE 206L – Fundamentals of Database Systems",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=Pt(24))

add_para("Presented to the Faculty of the", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para("College of Information Technology", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para("Franciscan College of the Immaculate Conception,",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para("Baybay City, Leyte Incorporated",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))

add_para("by", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=Pt(12))
add_para("[LASTNAME, Firstname M.]",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=Pt(36))

add_para("APRIL 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# APPROVAL SHEET
# ════════════════════════════════════════════════════════════════════════════

add_chapter_title("APPROVAL SHEET")

add_para(
    'This project titled "UPPP: A Community-Based Link Sharing Platform" '
    "prepared and submitted by [MEMBER'S NAME] in partial fulfillment of the "
    "requirements for the course ITE 206L: Fundamentals of Database Systems has "
    "been recommended for acceptance and approval for ORAL PRESENTATION.",
    space_after=Pt(24)
)

add_para("JILL ZARAH B. SABANDO, MIT",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=Pt(0))
add_para("Instructor, ITE 206L",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))

add_para("EVALUATOR", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=Pt(24))

for _ in range(3):
    add_para("_" * 40, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

add_para(" ", space_after=Pt(24))
add_para("PANEL OF EXAMINERS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=Pt(12))
add_para("Approved by the Committee on Oral Presentation with a grade of ______.",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

for _ in range(3):
    add_para("_" * 40, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

add_para(" ", space_after=Pt(24))
add_para(
    "ACCEPTED and APPROVED in partial fulfillment of the requirements for the course "
    "ITE 206L: Fundamentals of Database Systems.",
    space_after=Pt(12)
)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# BIOGRAPHICAL DATA
# ════════════════════════════════════════════════════════════════════════════

add_chapter_title("BIOGRAPHICAL DATA")

add_para(
    "[Insert your biographical data here. Include your full name, date of birth, "
    "address, educational background, and other relevant personal information as "
    "required by your instructor.]",
    italic=True
)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════

add_chapter_title("TABLE OF CONTENTS")

toc_entries = [
    ("Approval Sheet", "ii"),
    ("Biographical Data", "iii"),
    ("Table of Contents", "iv"),
    ("CHAPTER I – INTRODUCTION", ""),
    ("    Project Context", "1"),
    ("    Objectives", "2"),
    ("    Description of the Proposed System", "3"),
    ("    Scope and Limitation", "4"),
    ("CHAPTER II – DATABASE DESIGN", ""),
    ("    Entity Relationship Diagram", "5"),
    ("    Data Schema", "6"),
    ("    Database Dictionary", "7"),
    ("CHAPTER III – CRUD IMPLEMENTATION", ""),
    ("    Create", "10"),
    ("    Read (Display)", "12"),
    ("    Update (Edit)", "14"),
    ("    Delete (Remove)", "16"),
    ("    Security Measures", "17"),
    ("CHAPTER IV – USER MANUAL", ""),
    ("    Step-by-Step Guide", "19"),
    ("REFERENCES", "28"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment          = WD_ALIGN_PARAGRAPH.LEFT
    fmt.line_spacing_rule  = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before       = Pt(0)
    fmt.space_after        = Pt(3)
    is_chapter = entry.startswith("CHAPTER") or entry in (
        "Approval Sheet", "Biographical Data", "Table of Contents", "REFERENCES"
    )
    run = p.add_run(entry)
    set_run_font(run, bold=is_chapter)
    if page:
        tab = p.add_run(f"\t{page}")
        set_run_font(tab, bold=is_chapter)
        fmt.tab_stops.add_tab_stop(Inches(5.5), WD_ALIGN_PARAGRAPH.RIGHT)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER I – INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════

add_chapter_title("CHAPTER I")
add_chapter_title("INTRODUCTION")

# ── Project Context ──────────────────────────────────────────────────────────
add_section_heading("Project Context")

add_para(
    "The internet has grown into a massive collection of websites, tools, and "
    "resources. However, finding high-quality and useful links can be difficult "
    "because there is no single place where people can share and organize them "
    "together. Most search engines return too many results, making it hard to "
    "find what is truly useful and community-recommended."
)

add_para(
    "This project, called uppp, is a web-based community link-sharing platform "
    "that allows registered users to submit, discover, and vote on useful links. "
    "The system was developed as part of the requirements for ITE 206L – "
    "Fundamentals of Database Systems, with a focus on applying relational "
    "database concepts in a real-world web application."
)

add_para(
    "uppp uses PHP as its server-side programming language, MySQL as its "
    "database management system, and XAMPP as its local development server. "
    "The platform allows users to submit URLs along with a title and description, "
    "vote on submissions, leave comments, and organize content by category. "
    "An administrator panel is also included for content moderation and user management."
)

add_para(
    "The name 'uppp' is inspired by the concept of upvoting — a common feature "
    "in community-driven platforms where users decide what content is valuable "
    "by voting it up or down. The platform promotes a collaborative approach to "
    "content curation, where the community itself determines what is worth sharing."
)

# ── Objectives ───────────────────────────────────────────────────────────────
add_section_heading("Objectives")

add_para("The following are the main objectives of this project:")

objectives = [
    "To design and implement a normalized relational database that supports "
    "a community link-sharing platform.",
    "To apply CRUD (Create, Read, Update, Delete) operations using PHP and MySQL "
    "with PDO prepared statements.",
    "To implement user authentication, session management, and role-based access "
    "control for regular users and administrators.",
    "To apply security best practices including input validation, CSRF protection, "
    "password hashing, and XSS prevention.",
    "To build a functional content moderation system that allows administrators "
    "to manage submissions, users, reports, and categories.",
    "To provide a user-friendly interface that allows community members to "
    "submit, discover, vote on, and discuss useful online resources.",
]

for obj in objectives:
    add_bullet(obj)

# ── Description of Proposed System ───────────────────────────────────────────
add_section_heading("Description of the Proposed System")

add_para(
    "uppp is a web application that runs on a local XAMPP server using Apache "
    "and MySQL. It is built with plain PHP and follows a page-based routing "
    "structure, where each feature is a standalone PHP file. There is no "
    "external framework involved, making the codebase simple and easy to understand."
)

add_para(
    "The system has two types of users: regular users and administrators. "
    "Regular users can register an account, log in, submit links, vote on "
    "submissions, leave comments, search for content, browse by category, "
    "edit their profile, and report inappropriate content. Administrators "
    "have access to a dedicated admin panel where they can approve or reject "
    "submitted links, manage user accounts, resolve content reports, and "
    "manage categories."
)

add_para(
    "Submitted links go through a moderation workflow. When a user submits a "
    "link, it is saved with a 'pending' status. It only becomes visible to the "
    "public after an administrator approves it. Rejected submissions are only "
    "visible to the submitter and the administrator. This ensures that all "
    "publicly visible content has been reviewed."
)

add_para(
    "The voting system allows users to upvote or downvote any approved "
    "submission. Votes are toggled — clicking the same vote again removes it, "
    "while clicking the opposite vote switches it. The homepage features a "
    "'Top of the Week' section that ranks recent submissions using a formula "
    "that combines the net vote score with a recency bonus that decreases over "
    "seven days."
)

add_para(
    "The system also includes a comment section on each submission's detail page, "
    "a search feature that looks through titles, descriptions, and URLs, and a "
    "profile page where users can view their submission history and update their "
    "bio, avatar, and website link."
)

# ── Scope and Limitation ─────────────────────────────────────────────────────
add_section_heading("Scope and Limitation")

add_sub_heading("Scope")

add_para("The following features are included in the system:")

scope_items = [
    "User registration, login, and logout with session management.",
    "Link submission with title, URL, description, and category.",
    "Admin moderation workflow: pending, approved, and rejected status for submissions.",
    "Upvoting and downvoting of approved submissions via AJAX.",
    "Commenting on submission detail pages.",
    "Category browsing and full-text search across titles, descriptions, and URLs.",
    "User profile pages with bio, avatar upload, and website link.",
    "Content reporting system for submissions and comments.",
    "Admin panel for managing submissions, users, categories, and reports.",
    "Pagination on all list pages.",
]

for item in scope_items:
    add_bullet(item)

add_sub_heading("Limitation")

add_para("The following features are not included in the current version of the system:")

limit_items = [
    "Email verification upon registration — accounts are activated immediately without email confirmation.",
    "Real-time features such as live notifications or live updates to vote counts without page refresh.",
    "Mobile application — the system is only accessible through a web browser.",
    "Deployment on a public server — the system runs only on a local XAMPP installation.",
    "Rate limiting — there is no protection against automated or repeated form submissions.",
    "Nested comments or reply threads — comments are flat and non-threaded.",
    "Password reset functionality via email.",
]

for item in limit_items:
    add_bullet(item)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER II – DATABASE DESIGN
# ════════════════════════════════════════════════════════════════════════════

add_chapter_title("CHAPTER II")
add_chapter_title("DATABASE DESIGN")

# ── ERD ──────────────────────────────────────────────────────────────────────
add_section_heading("Entity Relationship Diagram")

add_para(
    "The database of uppp is composed of six (6) entities: users, categories, "
    "submissions, votes, comments, and reports. The diagram below describes the "
    "relationships among these entities."
)

add_para("[Insert ERD Diagram here]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)

add_para(
    "The following describes the relationships between the entities:"
)

erd_items = [
    "A USER can submit many SUBMISSIONS (one-to-many). Each submission belongs to exactly one user.",
    "A SUBMISSION belongs to one CATEGORY. A category can have many submissions (one-to-many). "
    "A category cannot be deleted if submissions are assigned to it.",
    "A USER can cast one VOTE per SUBMISSION (many-to-many resolved through the votes table). "
    "Each vote record stores either an upvote (+1) or a downvote (-1).",
    "A USER can post many COMMENTS on a SUBMISSION (one-to-many). Each comment belongs to one user "
    "and one submission.",
    "A USER can file many REPORTS. Each report targets either one SUBMISSION or one COMMENT, "
    "but never both at the same time. This is enforced by a CHECK constraint in the database.",
]

for item in erd_items:
    add_bullet(item)

# ── Data Schema ───────────────────────────────────────────────────────────────
add_section_heading("Data Schema")

add_para(
    "The following is the complete SQL schema used to create the uppp database. "
    "The database uses the utf8mb4 character set and utf8mb4_unicode_ci collation "
    "to support a wide range of characters."
)

schema_lines = [
    "CREATE DATABASE IF NOT EXISTS uppp",
    "    CHARACTER SET utf8mb4",
    "    COLLATE utf8mb4_unicode_ci;",
    "USE uppp;",
    "",
    "-- USERS table",
    "CREATE TABLE users (",
    "    id            INT AUTO_INCREMENT PRIMARY KEY,",
    "    username      VARCHAR(50)  NOT NULL UNIQUE,",
    "    email         VARCHAR(255) NOT NULL UNIQUE,",
    "    password_hash VARCHAR(255) NOT NULL,",
    "    role          ENUM('user','admin') NOT NULL DEFAULT 'user',",
    "    is_banned     TINYINT(1)   NOT NULL DEFAULT 0,",
    "    bio           TEXT         NULL,",
    "    avatar_path   VARCHAR(255) NULL,",
    "    website_url   VARCHAR(255) NULL,",
    "    created_at    DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP",
    ") ENGINE=InnoDB;",
    "",
    "-- CATEGORIES table",
    "CREATE TABLE categories (",
    "    id          INT AUTO_INCREMENT PRIMARY KEY,",
    "    name        VARCHAR(100) NOT NULL UNIQUE,",
    "    slug        VARCHAR(100) NOT NULL UNIQUE,",
    "    description TEXT         NULL",
    ") ENGINE=InnoDB;",
    "",
    "-- SUBMISSIONS table",
    "CREATE TABLE submissions (",
    "    id          INT AUTO_INCREMENT PRIMARY KEY,",
    "    user_id     INT          NOT NULL,",
    "    category_id INT          NOT NULL,",
    "    title       VARCHAR(255) NOT NULL,",
    "    url         VARCHAR(500) NOT NULL,",
    "    description TEXT         NOT NULL,",
    "    status      ENUM('pending','approved','rejected') NOT NULL DEFAULT 'pending',",
    "    created_at  DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP,",
    "    FOREIGN KEY (user_id)     REFERENCES users(id)      ON DELETE CASCADE,",
    "    FOREIGN KEY (category_id) REFERENCES categories(id) ON DELETE RESTRICT",
    ") ENGINE=InnoDB;",
    "",
    "-- VOTES table",
    "CREATE TABLE votes (",
    "    id            INT AUTO_INCREMENT PRIMARY KEY,",
    "    user_id       INT      NOT NULL,",
    "    submission_id INT      NOT NULL,",
    "    vote_type     TINYINT  NOT NULL,",
    "    created_at    DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,",
    "    UNIQUE KEY unique_user_vote (user_id, submission_id),",
    "    FOREIGN KEY (user_id)       REFERENCES users(id)       ON DELETE CASCADE,",
    "    FOREIGN KEY (submission_id) REFERENCES submissions(id) ON DELETE CASCADE",
    ") ENGINE=InnoDB;",
    "",
    "-- COMMENTS table",
    "CREATE TABLE comments (",
    "    id            INT AUTO_INCREMENT PRIMARY KEY,",
    "    user_id       INT      NOT NULL,",
    "    submission_id INT      NOT NULL,",
    "    body          TEXT     NOT NULL,",
    "    created_at    DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,",
    "    FOREIGN KEY (user_id)       REFERENCES users(id)       ON DELETE CASCADE,",
    "    FOREIGN KEY (submission_id) REFERENCES submissions(id) ON DELETE CASCADE",
    ") ENGINE=InnoDB;",
    "",
    "-- REPORTS table",
    "CREATE TABLE reports (",
    "    id            INT AUTO_INCREMENT PRIMARY KEY,",
    "    user_id       INT      NOT NULL,",
    "    submission_id INT      NULL,",
    "    comment_id    INT      NULL,",
    "    reason        TEXT     NOT NULL,",
    "    status        ENUM('pending','resolved','dismissed') NOT NULL DEFAULT 'pending',",
    "    created_at    DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,",
    "    FOREIGN KEY (user_id)       REFERENCES users(id)       ON DELETE CASCADE,",
    "    FOREIGN KEY (submission_id) REFERENCES submissions(id) ON DELETE CASCADE,",
    "    FOREIGN KEY (comment_id)    REFERENCES comments(id)    ON DELETE CASCADE,",
    "    CONSTRAINT chk_report_target CHECK (",
    "        (submission_id IS NOT NULL AND comment_id IS NULL) OR",
    "        (submission_id IS NULL     AND comment_id IS NOT NULL)",
    "    )",
    ") ENGINE=InnoDB;",
]

add_code_block(schema_lines)

# ── Database Dictionary ───────────────────────────────────────────────────────
add_section_heading("Database Dictionary")

# Helper for dictionary tables
def add_dict_table(table_name, description, columns):
    """columns = list of (column_name, data_type, constraints, description)"""
    add_sub_heading(f"Table: {table_name}")
    add_para(description, space_after=Pt(4))

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl.rows[0].cells
    headers = ["Column Name", "Data Type", "Constraints", "Description"]
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, bold=True, size=Pt(10))

    for col_name, dtype, constraints, desc in columns:
        row = tbl.add_row().cells
        data = [col_name, dtype, constraints, desc]
        for i, val in enumerate(data):
            p = row[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            set_run_font(run, size=Pt(10))

    # space after table
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(10)
    gap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


add_dict_table(
    "users",
    "Stores all registered user accounts including their credentials, role, "
    "and profile information.",
    [
        ("id",            "INT",          "PK, AUTO_INCREMENT",           "Unique identifier for each user"),
        ("username",      "VARCHAR(50)",  "NOT NULL, UNIQUE",             "Chosen display name of the user"),
        ("email",         "VARCHAR(255)", "NOT NULL, UNIQUE",             "Email address used for login"),
        ("password_hash", "VARCHAR(255)", "NOT NULL",                     "Bcrypt-hashed password"),
        ("role",          "ENUM",         "DEFAULT 'user'",               "Access level: 'user' or 'admin'"),
        ("is_banned",     "TINYINT(1)",   "NOT NULL, DEFAULT 0",          "1 if account is banned, 0 if active"),
        ("bio",           "TEXT",         "NULL",                         "Optional personal description"),
        ("avatar_path",   "VARCHAR(255)", "NULL",                         "File path of uploaded avatar image"),
        ("website_url",   "VARCHAR(255)", "NULL",                         "Optional personal website URL"),
        ("created_at",    "DATETIME",     "DEFAULT CURRENT_TIMESTAMP",    "Timestamp when the account was created"),
    ]
)

add_dict_table(
    "categories",
    "Stores the predefined content categories used to organize submissions.",
    [
        ("id",          "INT",          "PK, AUTO_INCREMENT", "Unique identifier for each category"),
        ("name",        "VARCHAR(100)", "NOT NULL, UNIQUE",   "Display name of the category"),
        ("slug",        "VARCHAR(100)", "NOT NULL, UNIQUE",   "URL-friendly version of the category name"),
        ("description", "TEXT",         "NULL",               "Short description of what the category covers"),
    ]
)

add_dict_table(
    "submissions",
    "Stores all link submissions made by users, along with their moderation status.",
    [
        ("id",          "INT",          "PK, AUTO_INCREMENT",        "Unique identifier for each submission"),
        ("user_id",     "INT",          "FK → users.id, CASCADE",    "The user who submitted the link"),
        ("category_id", "INT",          "FK → categories.id, RESTRICT", "The category assigned to the submission"),
        ("title",       "VARCHAR(255)", "NOT NULL",                  "Title or name of the submitted link"),
        ("url",         "VARCHAR(500)", "NOT NULL",                  "The full URL being submitted"),
        ("description", "TEXT",         "NOT NULL",                  "Short description of the link's content"),
        ("status",      "ENUM",         "DEFAULT 'pending'",         "Moderation status: pending, approved, or rejected"),
        ("created_at",  "DATETIME",     "DEFAULT CURRENT_TIMESTAMP", "Timestamp when the submission was created"),
    ]
)

add_dict_table(
    "votes",
    "Records each user's vote on a submission. Each user can only have one vote per submission.",
    [
        ("id",            "INT",      "PK, AUTO_INCREMENT",              "Unique identifier for each vote record"),
        ("user_id",       "INT",      "FK → users.id, CASCADE",          "The user who cast the vote"),
        ("submission_id", "INT",      "FK → submissions.id, CASCADE",    "The submission being voted on"),
        ("vote_type",     "TINYINT",  "NOT NULL",                        "+1 for upvote, -1 for downvote"),
        ("created_at",    "DATETIME", "DEFAULT CURRENT_TIMESTAMP",       "Timestamp when the vote was recorded"),
    ]
)

add_dict_table(
    "comments",
    "Stores user comments posted on submission detail pages.",
    [
        ("id",            "INT",      "PK, AUTO_INCREMENT",           "Unique identifier for each comment"),
        ("user_id",       "INT",      "FK → users.id, CASCADE",       "The user who wrote the comment"),
        ("submission_id", "INT",      "FK → submissions.id, CASCADE", "The submission being commented on"),
        ("body",          "TEXT",     "NOT NULL",                     "The text content of the comment"),
        ("created_at",    "DATETIME", "DEFAULT CURRENT_TIMESTAMP",    "Timestamp when the comment was posted"),
    ]
)

add_dict_table(
    "reports",
    "Stores content reports filed by users against submissions or comments.",
    [
        ("id",            "INT",      "PK, AUTO_INCREMENT",           "Unique identifier for each report"),
        ("user_id",       "INT",      "FK → users.id, CASCADE",       "The user who filed the report"),
        ("submission_id", "INT",      "FK → submissions.id, CASCADE, NULL", "The reported submission (NULL if reporting a comment)"),
        ("comment_id",    "INT",      "FK → comments.id, CASCADE, NULL",   "The reported comment (NULL if reporting a submission)"),
        ("reason",        "TEXT",     "NOT NULL",                     "Explanation of why the content was reported"),
        ("status",        "ENUM",     "DEFAULT 'pending'",            "Report status: pending, resolved, or dismissed"),
        ("created_at",    "DATETIME", "DEFAULT CURRENT_TIMESTAMP",    "Timestamp when the report was submitted"),
    ]
)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER III – CRUD IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════

add_chapter_title("CHAPTER III")
add_chapter_title("CRUD IMPLEMENTATION")

add_para(
    "This chapter presents the CRUD (Create, Read, Update, Delete) operations "
    "implemented in the uppp system. Each operation is demonstrated with actual "
    "PHP code from the project files, followed by an explanation of how the "
    "operation works."
)

# ── CREATE ────────────────────────────────────────────────────────────────────
add_section_heading("Create")

add_sub_heading("1. User Registration (register.php)")

add_para(
    "When a new user fills out the registration form and submits it, the system "
    "validates the input and inserts a new record into the users table. The "
    "password is never stored in plain text — it is hashed using PHP's built-in "
    "password_hash() function before being saved."
)

add_code_block([
    "// Collect and trim input",
    "$username = trim($_POST['username'] ?? '');",
    "$email    = trim($_POST['email']    ?? '');",
    "$password = $_POST['password']      ?? '';",
    "$confirm  = $_POST['confirm']       ?? '';",
    "",
    "// Validation",
    "if (strlen($username) < 3 || strlen($username) > 50)",
    "    $errors[] = 'Username must be 3–50 characters.';",
    "if (!filter_var($email, FILTER_VALIDATE_EMAIL))",
    "    $errors[] = 'Please enter a valid email address.';",
    "if (strlen($password) < 6)",
    "    $errors[] = 'Password must be at least 6 characters.';",
    "if ($password !== $confirm)",
    "    $errors[] = 'Passwords do not match.';",
    "",
    "// Check for duplicate username or email",
    "$stmt = $pdo->prepare(",
    "    'SELECT id FROM users WHERE username = ? OR email = ?'",
    ");",
    "$stmt->execute([$username, $email]);",
    "if ($stmt->fetch()) $errors[] = 'Username or email already taken.';",
    "",
    "// If no errors, insert new user",
    "if (empty($errors)) {",
    "    $hash = password_hash($password, PASSWORD_DEFAULT);",
    "    $stmt = $pdo->prepare(",
    "        'INSERT INTO users (username, email, password_hash)",
    "         VALUES (?, ?, ?)'",
    "    );",
    "    $stmt->execute([$username, $email, $hash]);",
    "    login_user($pdo, $email, $password);",
    "    redirect('index.php');",
    "}",
])

add_para(
    "The code above shows how the system collects the user's input, validates "
    "each field, checks for duplicate accounts in the database, and inserts the "
    "new user record with a hashed password. After successful registration, the "
    "user is automatically logged in and redirected to the homepage."
)

add_sub_heading("2. Link Submission (submit.php)")

add_para(
    "Authenticated users can submit a new link by filling out the submission form. "
    "The system validates all fields and checks if the URL has already been "
    "approved before inserting the record with a 'pending' status."
)

add_code_block([
    "// Collect input",
    "$title       = trim($_POST['title']       ?? '');",
    "$url         = trim($_POST['url']         ?? '');",
    "$description = trim($_POST['description'] ?? '');",
    "$category_id = (int)($_POST['category_id'] ?? 0);",
    "",
    "// Validate",
    "if (strlen($title) < 3 || strlen($title) > 255)",
    "    $errors[] = 'Title must be 3–255 characters.';",
    "if (!filter_var($url, FILTER_VALIDATE_URL) || strlen($url) > 500)",
    "    $errors[] = 'Please enter a valid URL (max 500 characters).';",
    "if (strlen($description) < 10)",
    "    $errors[] = 'Description must be at least 10 characters.';",
    "if ($category_id <= 0)",
    "    $errors[] = 'Please select a category.';",
    "",
    "// Check for duplicate approved URL",
    "$stmt = $pdo->prepare(",
    "    'SELECT id FROM submissions WHERE url = ? AND status = ?'",
    ");",
    "$stmt->execute([$url, 'approved']);",
    "if ($stmt->fetch()) $errors[] = 'This URL has already been submitted.';",
    "",
    "// Insert with pending status",
    "if (empty($errors)) {",
    "    $stmt = $pdo->prepare(",
    "        'INSERT INTO submissions",
    "             (user_id, category_id, title, url, description)",
    "         VALUES (?, ?, ?, ?, ?)'",
    "    );",
    "    $stmt->execute([",
    "        current_user_id(), $category_id, $title, $url, $description",
    "    ]);",
    "    flash('success', 'Submission received! Pending review.');",
    "    redirect('index.php');",
    "}",
])

add_para(
    "New submissions are always saved with a 'pending' status. They will only "
    "appear on the public homepage and category pages after an administrator "
    "approves them. The submitter can see their own pending submissions "
    "on their profile page."
)

# ── READ ──────────────────────────────────────────────────────────────────────
add_section_heading("Read (Display)")

add_sub_heading("1. Homepage with Top of the Week Ranking (index.php)")

add_para(
    "The homepage displays two lists: the 'Top of the Week' submissions and "
    "the most recent approved submissions. The top submissions are ranked using "
    "a formula that adds the net vote score to a recency bonus. The recency "
    "bonus starts at 7 and decreases by 1 for each day since the submission "
    "was created, stopping at 0 after seven days."
)

add_code_block([
    "// Top of the Week: approved submissions from the past 7 days",
    "// ranked by votes + recency bonus",
    "$topSql = '",
    "    SELECT",
    "        s.*,",
    "        u.username,",
    "        c.name AS category_name,",
    "        COALESCE(SUM(v.vote_type), 0) AS net_votes,",
    "        COALESCE(SUM(v.vote_type), 0)",
    "            + GREATEST(0, 7 - DATEDIFF(NOW(), s.created_at)) AS rank_score,",
    "        (SELECT COUNT(*) FROM comments cm",
    "         WHERE cm.submission_id = s.id) AS comment_count",
    "    FROM submissions s",
    "    JOIN  users u      ON s.user_id     = u.id",
    "    JOIN  categories c ON s.category_id = c.id",
    "    LEFT JOIN votes v  ON v.submission_id = s.id",
    "    WHERE s.status = ? AND s.created_at >= NOW() - INTERVAL 7 DAY",
    "    GROUP BY s.id",
    "    ORDER BY rank_score DESC",
    "    LIMIT 5",
    "';",
    "$stmt = $pdo->prepare($topSql);",
    "$stmt->execute(['approved']);",
    "$topSubmissions = $stmt->fetchAll();",
])

add_para(
    "The query uses JOIN operations to combine submission data with the "
    "corresponding user and category records. A LEFT JOIN on the votes table "
    "allows the query to count votes even when a submission has no votes yet. "
    "The COALESCE function returns 0 if there are no votes, preventing NULL "
    "values in the result. The GREATEST function ensures the recency bonus does "
    "not go below zero."
)

add_sub_heading("2. Search (search.php)")

add_para(
    "The search feature allows users to find approved submissions by entering "
    "keywords. The system searches across the title, description, and URL fields "
    "using the SQL LIKE operator."
)

add_code_block([
    "$q    = trim($_GET['q'] ?? '');",
    "$term = '%' . escape_like($q) . '%';",
    "",
    "$countSql = 'SELECT COUNT(*) FROM submissions s",
    "             JOIN users u ON s.user_id = u.id",
    "             WHERE s.status = ? AND (",
    "                 s.title       LIKE ? OR",
    "                 s.description LIKE ? OR",
    "                 s.url         LIKE ?",
    "             )';",
    "",
    "$pagination = paginate($pdo, $countSql,",
    "    ['approved', $term, $term, $term]);",
    "",
    "$stmt = $pdo->prepare(",
    "    'SELECT s.*, u.username, c.name AS category_name",
    "     FROM submissions s",
    "     JOIN users u      ON s.user_id     = u.id",
    "     JOIN categories c ON s.category_id = c.id",
    "     WHERE s.status = ? AND (",
    "         s.title LIKE ? OR s.description LIKE ? OR s.url LIKE ?",
    "     )",
    "     ORDER BY s.created_at DESC",
    "     LIMIT :limit OFFSET :offset'",
    ");",
    "$stmt->execute(['approved', $term, $term, $term]);",
    "$results = $stmt->fetchAll();",
])

add_para(
    "The escape_like() function escapes special characters (% and _) in the "
    "search term before passing it to the LIKE query. This prevents users from "
    "accidentally or intentionally manipulating the search pattern. Pagination "
    "is applied so that large result sets are split across multiple pages."
)

# ── UPDATE ────────────────────────────────────────────────────────────────────
add_section_heading("Update (Edit)")

add_sub_heading("1. Vote Toggle (vote.php)")

add_para(
    "The voting system handles three possible scenarios depending on whether "
    "the user has already voted. If the user has not voted yet, a new vote "
    "record is inserted. If the user clicks the same vote type again, the "
    "vote is deleted (toggled off). If the user clicks the opposite vote, "
    "the existing vote is updated. The result is returned as a JSON response "
    "to the JavaScript that sent the request."
)

add_code_block([
    "// Check if user already voted on this submission",
    "$stmt = $pdo->prepare(",
    "    'SELECT vote_type FROM votes",
    "     WHERE user_id = ? AND submission_id = ?'",
    ");",
    "$stmt->execute([$userId, $submissionId]);",
    "$existing = $stmt->fetchColumn();",
    "",
    "if ($existing === false) {",
    "    // No existing vote — insert",
    "    $stmt = $pdo->prepare(",
    "        'INSERT INTO votes (user_id, submission_id, vote_type)",
    "         VALUES (?, ?, ?)'",
    "    );",
    "    $stmt->execute([$userId, $submissionId, $voteType]);",
    "} elseif ((int)$existing === $voteType) {",
    "    // Same vote clicked again — delete (toggle off)",
    "    $stmt = $pdo->prepare(",
    "        'DELETE FROM votes",
    "         WHERE user_id = ? AND submission_id = ?'",
    "    );",
    "    $stmt->execute([$userId, $submissionId]);",
    "} else {",
    "    // Opposite vote — update",
    "    $stmt = $pdo->prepare(",
    "        'UPDATE votes SET vote_type = ?",
    "         WHERE user_id = ? AND submission_id = ?'",
    "    );",
    "    $stmt->execute([$voteType, $userId, $submissionId]);",
    "}",
    "",
    "// Return updated score as JSON",
    "header('Content-Type: application/json');",
    "echo json_encode(['score' => get_vote_score($pdo, $submissionId)]);",
])

add_para(
    "This approach ensures that the votes table always reflects the user's "
    "current voting intention. The UNIQUE constraint on (user_id, submission_id) "
    "in the database also prevents duplicate vote records at the database level."
)

add_sub_heading("2. Admin: Approve or Reject a Submission (admin/submissions.php)")

add_para(
    "Administrators can change the status of a pending submission to either "
    "'approved' or 'rejected'. The action is performed through a form POST "
    "with CSRF verification."
)

add_code_block([
    "if ($_SERVER['REQUEST_METHOD'] === 'POST' && verify_csrf()) {",
    "    $subId  = (int)($_POST['submission_id'] ?? 0);",
    "    $action = $_POST['action'] ?? '';",
    "",
    "    if ($subId > 0 && in_array($action, ['approved', 'rejected'])) {",
    "        $stmt = $pdo->prepare(",
    "            'UPDATE submissions SET status = ? WHERE id = ?'",
    "        );",
    "        $stmt->execute([$action, $subId]);",
    "        flash('success', 'Submission ' . $action . '.');",
    "    }",
    "    redirect('submissions.php?status=' . ($_GET['status'] ?? 'pending'));",
    "}",
])

add_para(
    "The in_array() check ensures that only the two allowed status values can "
    "be set through this form. This prevents an attacker from setting an "
    "arbitrary status value even if they bypass the form. The CSRF token is "
    "verified before any database change is made."
)

# ── DELETE ────────────────────────────────────────────────────────────────────
add_section_heading("Delete (Remove)")

add_sub_heading("Comment Deletion (submission.php)")

add_para(
    "Users can delete their own comments, and administrators can delete any "
    "comment. Before the deletion is performed, the system checks whether the "
    "currently logged-in user is the comment's author or an administrator."
)

add_code_block([
    "if (isset($_GET['delete_comment'])) {",
    "    require_login();",
    "    $commentId = (int)$_GET['delete_comment'];",
    "",
    "    // Fetch the comment record",
    "    $stmt = $pdo->prepare('SELECT * FROM comments WHERE id = ?');",
    "    $stmt->execute([$commentId]);",
    "    $comment = $stmt->fetch();",
    "",
    "    // Check authorization: must be comment owner or admin",
    "    if ($comment && (",
    "        current_user_id() === (int)$comment['user_id'] || is_admin()",
    "    )) {",
    "        $stmt = $pdo->prepare('DELETE FROM comments WHERE id = ?');",
    "        $stmt->execute([$commentId]);",
    "        flash('success', 'Comment deleted.');",
    "    }",
    "",
    "    redirect('submission.php?id=' . $id);",
    "}",
])

add_para(
    "The deletion only proceeds if the comment record actually exists and the "
    "authorization check passes. After the deletion, the user is redirected "
    "back to the same submission page with a success message. Because comments "
    "use ON DELETE CASCADE in the database, any associated reports pointing "
    "to this comment are also automatically removed."
)

# ── SECURITY ──────────────────────────────────────────────────────────────────
add_section_heading("Security Measures")

add_sub_heading("Input Validation")

add_para(
    "All user input is validated on the server side before being processed. "
    "PHP's built-in functions are used to validate data types and formats:"
)

add_code_block([
    "// Length checks",
    "if (strlen($username) < 3 || strlen($username) > 50) ...",
    "",
    "// Email format",
    "if (!filter_var($email, FILTER_VALIDATE_EMAIL)) ...",
    "",
    "// URL format",
    "if (!filter_var($url, FILTER_VALIDATE_URL)) ...",
    "",
    "// Integer cast for IDs (prevents SQL injection through type mismatch)",
    "$id = (int)($_GET['id'] ?? 0);",
    "if ($id <= 0) { redirect('index.php'); }",
])

add_sub_heading("Prepared Statements")

add_para(
    "All database queries in uppp use PDO prepared statements. This means user "
    "input is never directly embedded in SQL strings, making SQL injection "
    "attacks impossible."
)

add_code_block([
    "// UNSAFE (never done in uppp):",
    "// $pdo->query(\"SELECT * FROM users WHERE email = '$email'\");",
    "",
    "// SAFE (always done this way):",
    "$stmt = $pdo->prepare('SELECT * FROM users WHERE email = ?');",
    "$stmt->execute([$email]);",
    "$user = $stmt->fetch();",
])

add_para(
    "The PDO connection is also configured with emulated prepares turned off "
    "(PDO::ATTR_EMULATE_PREPARES set to false), which ensures that real prepared "
    "statements are used at the MySQL level."
)

add_sub_heading("Error Handling")

add_para(
    "PDO is configured to throw exceptions on errors (PDO::ERRMODE_EXCEPTION). "
    "All database operations are wrapped in a way that allows errors to surface "
    "clearly during development. Output is always escaped using the sanitize() "
    "function before being displayed in HTML:"
)

add_code_block([
    "// sanitize() in includes/functions.php",
    "function sanitize(string $value): string {",
    "    return htmlspecialchars(trim($value), ENT_QUOTES, 'UTF-8');",
    "}",
    "",
    "// Usage in views:",
    "echo sanitize($submission['title']);",
    "echo sanitize($user['username']);",
])

add_para(
    "Additionally, all state-changing forms include a CSRF token that is "
    "verified on the server before any data is modified. This prevents "
    "Cross-Site Request Forgery (CSRF) attacks where an attacker tricks a "
    "logged-in user into submitting an unintended form action."
)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER IV – USER MANUAL
# ════════════════════════════════════════════════════════════════════════════

add_chapter_title("CHAPTER IV")
add_chapter_title("USER MANUAL")

add_para(
    "This chapter provides a step-by-step guide on how to use the uppp system. "
    "The manual covers both regular user and administrator functions. "
    "Screenshots should be inserted at the designated areas."
)

manual_steps = [
    (
        "Step 1: Accessing the Homepage",
        "Open a web browser and go to http://localhost/uppp/. The homepage will "
        "display the 'Top of the Week' section at the top, which shows the five "
        "highest-ranked submissions from the past seven days. Below it, the most "
        "recently approved submissions are listed with pagination controls at the bottom.",
    ),
    (
        "Step 2: Registering an Account",
        "Click the 'Register' link in the top navigation bar. Fill out the registration "
        "form by entering a username (3–50 characters), a valid email address, and a "
        "password (minimum 6 characters). Enter the password again in the 'Confirm "
        "Password' field. Click the 'Register' button. If all inputs are valid and the "
        "username and email are not already taken, the account is created and the user "
        "is automatically logged in.",
    ),
    (
        "Step 3: Logging In",
        "Click the 'Login' link in the navigation bar. Enter the registered email "
        "address and password. Click the 'Login' button. If the credentials are "
        "correct and the account is not banned, the user is redirected to the "
        "homepage with a success message. A failed login will show an error message "
        "without revealing which specific field was incorrect.",
    ),
    (
        "Step 4: Submitting a Link",
        "After logging in, click the 'Submit' link in the navigation bar. Fill out "
        "the submission form with a title (3–255 characters), the URL to share, a "
        "description of at least 10 characters, and a category. Click 'Submit'. "
        "The link is saved with a 'pending' status and will only appear publicly "
        "after an administrator approves it. A notification message confirms that "
        "the submission was received.",
    ),
    (
        "Step 5: Voting on a Submission",
        "On the homepage or any submission list, click the upvote (▲) or downvote "
        "(▼) button beside a submission. The vote score updates instantly without "
        "reloading the page. Clicking the same vote again removes the vote. Clicking "
        "the opposite vote switches the direction. Users must be logged in to vote.",
    ),
    (
        "Step 6: Viewing and Leaving a Comment",
        "Click the title of any approved submission to open its detail page. The "
        "full description, vote score, and comment section will be displayed. "
        "To leave a comment, type in the text area at the bottom of the comment "
        "section and click 'Post Comment'. The comment appears immediately on the page.",
    ),
    (
        "Step 7: Searching for Content",
        "Click on the search bar in the navigation bar and type any keyword. Press "
        "Enter or click the search button. The system will display approved submissions "
        "whose title, description, or URL contains the keyword. Results are paginated "
        "if there are many matches.",
    ),
    (
        "Step 8: Browsing by Category",
        "Click on any category name in the navigation bar. The category page will "
        "display all approved submissions that belong to that category, sorted by "
        "newest first, with pagination.",
    ),
    (
        "Step 9: Editing a Profile",
        "Click on your username in the navigation bar, then click 'Edit Profile'. "
        "Users can update their bio and website URL. To upload a profile avatar, "
        "click 'Choose File' and select an image file (JPG, PNG, or GIF, maximum "
        "2MB). Click 'Save Changes'. The avatar is resized automatically and the "
        "old one is replaced.",
    ),
    (
        "Step 10: Reporting Content",
        "On any submission's detail page, click the 'Report' link. On any comment, "
        "click the 'Report' link beside it. Fill in the reason for the report and "
        "submit the form. Reports are sent to the administrator for review. Users "
        "must be logged in to submit a report.",
    ),
    (
        "Step 11: Admin — Logging In as Administrator",
        "Log in using the administrator credentials (email: admin@uppp.local, "
        "password: admin123 for the default seeded account). After logging in, "
        "an 'Admin' link will appear in the navigation bar. Click it to access "
        "the admin dashboard.",
    ),
    (
        "Step 12: Admin — Approving or Rejecting Submissions",
        "In the admin panel, click 'Submissions' in the sidebar. Pending "
        "submissions are shown by default. Review each submission and click "
        "'Approve' to make it publicly visible or 'Reject' to hide it. Approved "
        "and rejected submissions can also be viewed using the filter tabs.",
    ),
    (
        "Step 13: Admin — Managing Users",
        "Click 'Users' in the admin sidebar to see all registered accounts. "
        "Administrators can ban a user (preventing them from logging in), "
        "unban a previously banned user, promote a user to admin, or demote "
        "an admin back to a regular user.",
    ),
    (
        "Step 14: Admin — Resolving Reports",
        "Click 'Reports' in the admin sidebar to view all pending content reports. "
        "Each report shows the reporter's name, the type of content (submission or "
        "comment), and the reason given. Administrators can mark a report as "
        "'Resolved' (action was taken) or 'Dismissed' (report was not valid).",
    ),
    (
        "Step 15: Admin — Managing Categories",
        "Click 'Categories' in the admin sidebar to view, add, edit, or delete "
        "content categories. To add a new category, fill in the name and optional "
        "description and click 'Add Category'. The slug is generated automatically. "
        "A category cannot be deleted if there are submissions assigned to it — "
        "those submissions must be reassigned or deleted first.",
    ),
]

for i, (step_title, step_body) in enumerate(manual_steps, 1):
    add_sub_heading(step_title)
    add_para(step_body)
    add_para("[Screenshot placeholder]",
             align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=Pt(10))

page_break()


# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════

add_chapter_title("REFERENCES")

references = [
    "Gilmore, W. J. (2010). Beginning PHP and MySQL: From novice to professional "
    "(4th ed.). Apress.",

    "Kroenke, D. M., & Auer, D. J. (2015). Database concepts (7th ed.). Pearson "
    "Education.",

    "MySQL AB. (2023). MySQL 8.0 reference manual. Oracle Corporation. "
    "https://dev.mysql.com/doc/refman/8.0/en/",

    "The PHP Group. (2024). PHP: Hypertext preprocessor documentation. "
    "https://www.php.net/manual/en/",

    "The PHP Group. (2024). PHP data objects (PDO). "
    "https://www.php.net/manual/en/book.pdo.php",

    "OWASP Foundation. (2021). OWASP top ten. Open Web Application Security Project. "
    "https://owasp.org/www-project-top-ten/",

    "OWASP Foundation. (2023). Cross-site request forgery (CSRF) prevention cheat sheet. "
    "https://cheatsheetseries.owasp.org/cheatsheets/Cross-Site_Request_Forgery_Prevention_Cheat_Sheet.html",

    "OWASP Foundation. (2023). SQL injection prevention cheat sheet. "
    "https://cheatsheetseries.owasp.org/cheatsheets/SQL_Injection_Prevention_Cheat_Sheet.html",

    "Apache Friends. (2024). XAMPP — Apache + MariaDB + PHP + Perl. "
    "https://www.apachefriends.org/",

    "Codd, E. F. (1970). A relational model of data for large shared data banks. "
    "Communications of the ACM, 13(6), 377–387. https://doi.org/10.1145/362384.362685",
]

for ref in references:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.line_spacing_rule  = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before       = Pt(0)
    fmt.space_after        = Pt(6)
    fmt.left_indent        = Inches(0.5)
    fmt.first_line_indent  = Inches(-0.5)   # hanging indent
    run = p.add_run(ref)
    set_run_font(run)


# ════════════════════════════════════════════════════════════════════════════
# Save
# ════════════════════════════════════════════════════════════════════════════

output_path = "uppp_documentation.docx"
doc.save(output_path)
print(f"Documentation saved to: {output_path}")
